"""
service.py
──────────
Pure document-generation logic, extracted from the Tkinter SLA Generator.
No GUI dependencies — takes a data dict, returns the filled .docx as bytes.
Also provides generate_and_store() and send_esign_request() which persist
documents to Supabase Storage and optionally dispatch Dropbox Sign requests.
"""

import copy, io, json, re, tempfile, os
from datetime import datetime, timedelta, timezone
from docx import Document
from supabase import create_client

import dropbox_sign
from dropbox_sign import ApiClient, Configuration, SignatureRequestApi, models as ds_models

from .config import (
    TEMPLATE_PATH, ORDINAL, MONTHS,
    SUPABASE_URL, SUPABASE_KEY,
    DROPBOX_SIGN_API_KEY,
)
from .models import SLAGenerateRequest, SLAEsignRequest

DOTS = r"[…\.]{2,}"


def _fmt_naira(value) -> str:
    try:
        n = float(str(value).replace(",", "").replace("₦", "").replace("N", "").strip())
        return f"₦{int(n):,}" if n == int(n) else f"₦{n:,.2f}"
    except Exception:
        return str(value)


def _set_cell(row, col_idx: int, text: str):
    cell = row.cells[col_idx]
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)


def generate_sla_bytes(req: SLAGenerateRequest) -> bytes:
    """
    Fill the SLA Word template with the supplied request data.
    Returns the completed document as raw bytes (suitable for a file download response).
    """
    doc   = Document(TEMPLATE_PATH)
    paras = doc.paragraphs

    data = req.model_dump()
    data["company_name"] = data["company_name"].strip().upper()

    # 1. Cover page — para 25
    p25 = paras[25]
    for r in p25.runs:
        r.text = ""
    if p25.runs:
        p25.runs[0].text = data["company_name"]

    # 2. Opening body paragraph — para 31
    p31 = paras[31]
    #  2a. Day + Month
    for run in p31.runs:
        if re.search(DOTS, run.text) and "day of" in run.text:
            run.text = re.sub(DOTS, ORDINAL[data["contract_day"]], run.text, count=1)
            run.text = re.sub(DOTS, data["contract_month"],        run.text, count=1)
            break
    #  2b. Company name (standalone dots run)
    for run in p31.runs:
        t = run.text.strip()
        if re.fullmatch(r"[…\.\s]+", t) or (
            re.search(DOTS, t) and len(t) < 60 and "address" not in t
        ):
            run.text = " " + data["company_name"] + " "
            break
    #  2c. Company address
    for run in p31.runs:
        if "address is at" in run.text:
            run.text = re.sub(DOTS, data["company_address"], run.text, count=1)
            break

    # 3. Premium paragraph — para 56
    p56  = paras[56]
    full = "".join(r.text for r in p56.runs)
    full = re.sub(DOTS, data["premium_naira"],     full, count=1)
    full = re.sub(DOTS, data["premium_words"],     full, count=1)
    full = re.sub(DOTS, data["num_beneficiaries"], full, count=1)
    for r in p56.runs:
        r.text = ""
    if p56.runs:
        p56.runs[0].text = full

    # 4. Plans table
    table    = doc.tables[0]
    data_row = table.rows[1]
    total_tr = table.rows[2]._tr

    plans        = data["plans"]
    total_lives  = sum(int(str(p.get("num_lives", 0)).replace(",", "") or 0) for p in plans)
    total_amount = sum(
        float(str(p.get("amount", 0)).replace(",", "").replace("₦", "") or 0) for p in plans
    )

    first = plans[0] if plans else {}
    _set_cell(data_row, 0, "1.")
    _set_cell(data_row, 1, first.get("plan_type", ""))
    _set_cell(data_row, 2, first.get("description", ""))
    _set_cell(data_row, 3, str(first.get("num_lives", "")))
    a = first.get("amount", "")
    _set_cell(data_row, 4, _fmt_naira(a) if a else "")

    for idx, plan in enumerate(plans[1:], start=2):
        new_tr = copy.deepcopy(data_row._tr)
        total_tr.addprevious(new_tr)
        new_row = table.rows[idx]
        _set_cell(new_row, 0, f"{idx}.")
        _set_cell(new_row, 1, plan.get("plan_type", ""))
        _set_cell(new_row, 2, plan.get("description", ""))
        _set_cell(new_row, 3, str(plan.get("num_lives", "")))
        aa = plan.get("amount", "")
        _set_cell(new_row, 4, _fmt_naira(aa) if aa else "")

    total_row = table.rows[-1]
    _set_cell(total_row, 2, f"TOTAL  ({total_lives:,} lives)")
    grand = int(total_amount) if total_amount == int(total_amount) else total_amount
    _set_cell(total_row, 4, _fmt_naira(grand))

    # 5. Period of cover — para 123
    p123 = paras[123]
    full = "".join(r.text for r in p123.runs)
    full = re.sub(DOTS, ORDINAL[data["start_day"]],                              full, count=1)
    full = re.sub(DOTS, data["start_month"],                                     full, count=1)
    full = re.sub(DOTS, f"{ORDINAL[data['end_day']]} {data['end_month']}",       full, count=1)
    full = full.replace("2026", data["start_year"], 1)
    full = full.replace("2027", data["end_year"],   1)
    for r in p123.runs:
        r.text = ""
    if p123.runs:
        p123.runs[0].text = full

    # 6. Signature page — para 268
    p268 = paras[268]
    for r in p268.runs:
        r.text = ""
    if p268.runs:
        p268.runs[0].text = data["company_name"]

    # Return as bytes (no temp file needed)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _supabase_client():
    return create_client(SUPABASE_URL, SUPABASE_KEY)


def _upload_docx(sb, docx_bytes: bytes, company_name: str, year: str) -> str:
    """Upload docx to Supabase Storage and return the storage path."""
    slug = re.sub(r"[^\w]", "_", company_name.strip().upper())[:40]
    ts   = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    path = f"{year}/{slug}_{ts}.docx"
    sb.storage.from_("sla-documents").upload(
        path,
        docx_bytes,
        {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    )
    return path


def _signed_url(sb, path: str) -> str:
    """Create a signed URL valid for ~10 years (3650 days)."""
    result = sb.storage.from_("sla-documents").create_signed_url(path, 3650 * 24 * 3600)
    return result.get("signedURL") or result.get("signed_url") or result["data"]["signedURL"]


def _month_num(month_name: str) -> int:
    return MONTHS.index(month_name) + 1


# ─────────────────────────────────────────────────────────────────────────────
# generate_and_store
# ─────────────────────────────────────────────────────────────────────────────

def generate_and_store(req: SLAGenerateRequest, generated_by: str = "") -> dict:
    """
    Generate the SLA document, upload to Supabase Storage, record in
    sla_documents table, and return a dict with download details.
    """
    docx_bytes = generate_sla_bytes(req)
    sb         = _supabase_client()

    storage_path = _upload_docx(sb, docx_bytes, req.company_name, req.start_year)
    download_url = _signed_url(sb, storage_path)

    now         = datetime.now(timezone.utc)
    expires_at  = now + timedelta(days=365)

    start_num = _month_num(req.start_month)
    end_num   = _month_num(req.end_month)

    row = {
        "company_name":           req.company_name.strip().upper(),
        "contract_start":         f"{req.start_year}-{start_num:02d}-{req.start_day:02d}",
        "contract_end":           f"{req.end_year}-{end_num:02d}-{req.end_day:02d}",
        "generated_by":           generated_by,
        "action":                 "download",
        "storage_path":           storage_path,
        "download_url":           download_url,
        "signature_request_id":   None,
        "test_mode":              None,
        "signers":                None,
        "created_at":             now.isoformat(),
        "expires_at":             expires_at.isoformat(),
    }

    result = sb.table("sla_documents").insert(row).execute()
    record = result.data[0] if result.data else {}

    return {
        "id":           record.get("id"),
        "company_name": row["company_name"],
        "download_url": download_url,
        "storage_path": storage_path,
    }


# ─────────────────────────────────────────────────────────────────────────────
# send_esign_request
# ─────────────────────────────────────────────────────────────────────────────

def send_esign_request(req: SLAEsignRequest, generated_by: str = "") -> dict:
    """
    Generate the SLA document, upload to Supabase Storage, dispatch a
    Dropbox Sign multi-signer request, and record in sla_documents.
    """
    docx_bytes = generate_sla_bytes(req)
    sb         = _supabase_client()

    storage_path = _upload_docx(sb, docx_bytes, req.company_name, req.start_year)
    download_url = _signed_url(sb, storage_path)

    # Write to temp file for Dropbox Sign upload
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        configuration = Configuration(username=DROPBOX_SIGN_API_KEY)

        signers = [
            ds_models.SubSignatureRequestSigner(
                email_address=req.director_email,
                name=req.director_name,
                order=0,
            ),
            ds_models.SubSignatureRequestSigner(
                email_address=req.legal_head_email,
                name=req.legal_head_name,
                order=1,
            ),
        ]
        if req.client_name and req.client_email and not req.test_mode:
            signers.append(ds_models.SubSignatureRequestSigner(
                email_address=req.client_email,
                name=req.client_name,
                order=2,
            ))

        send_request = ds_models.SignatureRequestSendRequest(
            test_mode=req.test_mode,
            title=f"Service Level Agreement — {req.company_name.strip().upper()}",
            subject=f"Please sign: Clearline SLA for {req.company_name.strip().upper()}",
            message=(
                f"Please review and sign the Service Level Agreement between "
                f"Clearline International Limited and {req.company_name.strip().upper()}. "
                "You can sign directly in your browser or download the document."
            ),
            signers=signers,
            cc_email_addresses=[req.hr_email] if req.hr_email else [],
            files=[open(tmp_path, "rb")],
            signing_options=ds_models.SubSigningOptions(
                draw=True, type=True, upload=True,
                phone=False, default_type="draw",
            ),
        )

        with ApiClient(configuration) as api_client:
            response = SignatureRequestApi(api_client).signature_request_send(send_request)

        sig_req              = response.signature_request
        signature_request_id = sig_req.signature_request_id
    finally:
        os.unlink(tmp_path)

    now        = datetime.now(timezone.utc)
    expires_at = now + timedelta(days=365)

    start_num = _month_num(req.start_month)
    end_num   = _month_num(req.end_month)

    signers_json = {
        "director_name":    req.director_name,
        "director_email":   req.director_email,
        "legal_head_name":  req.legal_head_name,
        "legal_head_email": req.legal_head_email,
        "hr_email":         req.hr_email,
        "client_name":      req.client_name,
        "client_email":     req.client_email,
    }

    row = {
        "company_name":           req.company_name.strip().upper(),
        "contract_start":         f"{req.start_year}-{start_num:02d}-{req.start_day:02d}",
        "contract_end":           f"{req.end_year}-{end_num:02d}-{req.end_day:02d}",
        "generated_by":           generated_by,
        "action":                 "esign",
        "storage_path":           storage_path,
        "download_url":           download_url,
        "signature_request_id":   signature_request_id,
        "test_mode":              req.test_mode,
        "signers":                json.dumps(signers_json),
        "created_at":             now.isoformat(),
        "expires_at":             expires_at.isoformat(),
    }

    result = sb.table("sla_documents").insert(row).execute()
    record = result.data[0] if result.data else {}

    return {
        "id":                    record.get("id"),
        "company_name":          row["company_name"],
        "download_url":          download_url,
        "storage_path":          storage_path,
        "signature_request_id":  signature_request_id,
        "test_mode":             req.test_mode,
    }
